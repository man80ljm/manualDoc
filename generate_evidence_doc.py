import os
import re
import logging
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
from tkinter import messagebox, filedialog
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.scrolled import ScrolledText
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

# Configure logging
logging.basicConfig(level=logging.INFO, filename='doc_generator.log', format='%(asctime)s - %(levelname)s - %(message)s')

# Supported file extensions
SUPPORTED_EXTENSIONS = {'.png', '.jpg', '.jpeg'}

# Global variable to store the selected folder
selected_folder = ""

def natural_sort_key(s):
    """Extract numbers from filename for natural sorting."""
    return [int(c) if c.isdigit() else c.lower() for c in re.split(r'(\d+)', s)]

def to_chinese_num(num):
    """Convert number to Chinese numeral with a comma."""
    chinese_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    if num <= 10:
        return f"{chinese_nums[num]}、"
    elif num < 20:
        return f"十{chinese_nums[num - 10]}、"
    else:
        return f"{num}、"

def convert_folder_name(folder_name, level):
    """Convert folder name like '1.美女' to '一、美女' for level 1, keep original for other levels."""
    match = re.match(r'(\d+)\.(.*)', folder_name)
    if match:
        num = int(match.group(1))
        rest = match.group(2).strip()
        if level == 1:
            return f"{to_chinese_num(num)}{rest}"
        else:
            return folder_name
    return folder_name

def check_unsupported_files(folder_path):
    """Check for unsupported file types in the folder and return a list of them."""
    unsupported_files = []
    for item in Path(folder_path).iterdir():
        if item.is_dir():
            unsupported_files.extend(check_unsupported_files(item))
        elif item.is_file() and item.suffix.lower() not in SUPPORTED_EXTENSIONS:
            unsupported_files.append(str(item))
    return unsupported_files

def rename_image(image, new_name, update_status):
    """Rename a single image file with error handling."""
    try:
        image.rename(new_name)
        update_status(f"Renamed {image} to {new_name}", "success")
        logging.info(f"Renamed {image} to {new_name}")
    except Exception as e:
        update_status(f"Error renaming {image}: {e}", "danger")
        logging.error(f"Error renaming {image}: {e}")

def rename_images(folder_path, update_status):
    """Recursively rename images in all folders to 1.jpg, 2.jpg, etc., using threads."""
    with ThreadPoolExecutor() as executor:
        for item in Path(folder_path).iterdir():
            if item.is_dir():
                rename_images(item, update_status)
            elif item.is_file() and item.suffix.lower() in SUPPORTED_EXTENSIONS:
                try:
                    images = [f for f in item.parent.iterdir() if f.suffix.lower() in SUPPORTED_EXTENSIONS]
                    images = sorted(images, key=lambda x: natural_sort_key(x.name))
                    for i, image in enumerate(images, 1):
                        new_name = item.parent / f"{i}{image.suffix.lower()}"
                        if image != new_name and not new_name.exists():
                            executor.submit(rename_image, image, new_name, update_status)
                except Exception as e:
                    update_status(f"Error processing folder {item.parent}: {e}", "danger")
                    logging.error(f"Error processing folder {item.parent}: {e}")

def collect_files(folder_path, level=1):
    """Recursively collect folder structure and files for any level, only including supported files."""
    structure = []
    for item in sorted(Path(folder_path).iterdir(), key=lambda x: natural_sort_key(x.name)):
        if item.is_dir():
            subfolders = collect_files(item, level + 1)
            files = []
            for file in sorted(item.iterdir(), key=lambda x: natural_sort_key(x.name)):
                if file.suffix.lower() in SUPPORTED_EXTENSIONS:
                    files.append(file)
            if subfolders or files:
                structure.append((item.name, files, subfolders))
    return structure

def count_total_images(structure):
    """Count total images in the structure for progress bar."""
    total = 0
    for folder_name, files, subfolders in structure:
        total += len(files)
        if subfolders:
            total += count_total_images(subfolders)
    return total

def get_image_dimensions(image_path):
    """Get the dimensions of an image using PIL with error handling."""
    try:
        with Image.open(image_path) as img:
            width, height = img.size
            return width, height
    except Exception as e:
        logging.error(f"Error getting image dimensions for {image_path}: {e}")
        return None, None

def create_document(structure, root_path, output_path, update_progress, update_status):
    """Create Word document with content pages, supporting multiple levels."""
    logging.info(f"Creating document at: {output_path}")
    doc = Document()
    
    # Set page margins for all pages
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    
    # Calculate available dimensions
    page_width = 21  # cm (A4 width)
    page_height = 29.7  # cm (A4 height)
    margin = 1.27  # cm
    available_width = page_width - 2 * margin
    available_width_in = available_width / 2.54
    available_height = page_height - 2 * margin
    available_height_in = available_height / 2.54
    
    # Estimate title height
    title_height_in = (10.5 / 72) * 1.5  # Adjusted for 1.5 line spacing
    scale_factor = 0.9  # Reduced for more margin
    
    # Customize Heading styles
    styles = doc.styles
    style1 = styles['Heading 1']
    style1.font.name = 'SimSun'
    style1._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style1.font.size = Pt(10.5)
    style1.font.bold = True
    style1.paragraph_format.space_before = Pt(0)
    style1.paragraph_format.space_after = Pt(0)
    style1.paragraph_format.line_spacing = 1.0
    style1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    style2 = styles['Heading 2']
    style2.font.name = 'SimSun'
    style2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style2.font.size = Pt(10.5)
    style2.font.bold = False
    style2.paragraph_format.space_before = Pt(0)
    style2.paragraph_format.space_after = Pt(0)
    style2.paragraph_format.line_spacing = 1.0
    style2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    total_images = count_total_images(structure)
    if total_images == 0:
        update_status("没有找到任何支持的图片文件！", "danger")
        logging.warning("No supported image files found")
        return False
    
    current_image = 0
    
    def add_headings(folder_structure, level=1, is_first_secondary=False):
        nonlocal current_image
        for i, (folder_name, files, subfolders) in enumerate(folder_structure, 1):
            display_name = convert_folder_name(folder_name, level)
            
            # Add heading
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(display_name)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(10.5)
            run.bold = (level == 1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.style = f'Heading {level}'
            paragraph.paragraph_format.keep_with_next = True
            
            # Add images
            for file in files:
                update_status(f"Adding image: {file}", "info")
                logging.info(f"Adding image: {file}")
                try:
                    img_width, img_height = get_image_dimensions(file)
                    if img_width is None or img_height is None:
                        update_status(f"Skipping {file}: Unable to read image dimensions", "danger")
                        continue
                    aspect_ratio = img_width / img_height
                    image_width = available_width_in
                    image_height = image_width / aspect_ratio
                    total_height = image_height + title_height_in
                    if total_height > available_height_in:
                        image_height = available_height_in - title_height_in
                        image_width = image_height * aspect_ratio
                        if image_width > available_width_in:
                            image_width = available_width_in
                            image_height = image_width / aspect_ratio
                    image_width *= scale_factor
                    image_height *= scale_factor
                    
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(file), width=Inches(image_width), height=Inches(image_height))
                    current_image += 1
                    update_progress(current_image / total_images * 100)
                except Exception as e:
                    update_status(f"Error adding image {file}: {e}", "danger")
                    logging.error(f"Error adding image {file}: {e}")
            
            # Add subfolders
            if subfolders:
                if level > 1 and not is_first_secondary:
                    doc.add_page_break()
                add_headings(subfolders, level + 1, is_first_secondary=(level == 1))
    
    add_headings(structure)
    
    output_path_str = str(output_path)
    try:
        if os.path.exists(output_path_str):
            with open(output_path_str, 'a') as f:
                pass
        doc.save(output_path_str)
        update_status("文件创建完成！", "success")
        logging.info(f"Document creation completed at: {output_path_str}")
        return True
    except IOError:
        messagebox.showwarning("警告", "请关闭文件再导出！")
        update_status("Failed to save: File is open or permission denied", "danger")
        logging.error("Failed to save: File is open or permission denied")
        return False

class DocGeneratorApp:
    def __init__(self):
        self.root = ttk.Window(themename="flatly")
        self.root.title("佐证材料生成")
        self.root.geometry("500x350")
        self.root.minsize(500, 350)
        self.root.resizable(True, True)
        try:
            self.root.iconbitmap('manualDoc.ico')
        except:
            logging.warning("Failed to load manualDoc.ico")
        
        self.selected_folder = ""
        self.setup_ui()
    
    def setup_ui(self):
        frame = ttk.Frame(self.root)
        frame.pack(pady=10, padx=10, fill="both", expand=True)
        
        self.label = ttk.Label(frame, text="没有选择文件夹!", bootstyle="info")
        self.label.pack(pady=10)
        
        # Create a centered container for buttons
        button_container = ttk.Frame(frame)
        button_container.pack(expand=True, anchor="center")
        
        self.btn_select = ttk.Button(button_container, text="加载文件夹", command=self.select_folder, bootstyle="primary")
        self.btn_select.pack(side="left", padx=5)
        self.btn_generate = ttk.Button(button_container, text="生成文件", command=self.generate_doc, bootstyle="success")
        self.btn_generate.pack(side="left", padx=5)
        
        self.progress_bar = ttk.Progressbar(frame, orient="horizontal", length=300, mode="determinate", bootstyle="striped")
        self.progress_bar.pack(pady=10, fill="x", padx=20)
        
        self.status_text = ScrolledText(frame, height=5, wrap="word", autohide=True)
        self.status_text.pack(pady=5, padx=10, fill="both")
        self.status_text.tag_config("success", foreground="green")
        self.status_text.tag_config("danger", foreground="red")
        self.status_text.tag_config("info", foreground="blue")
    
    def select_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.selected_folder = folder_path
            self.label.config(text=f"已选择: {folder_path}")
            self.update_status(f"已选择文件夹: {folder_path}", "info")
            logging.info(f"Folder selected: {self.selected_folder}")
        else:
            self.label.config(text="没有选择文件夹!")
            self.update_status("没有选择文件夹!", "danger")
            logging.warning("No folder selected")
    
    def update_status(self, text, style="info"):
        self.status_text.insert("end", f"{text}\n", style)
        self.status_text.see("end")
        self.root.update()
    
    def generate_doc(self):
        if not self.selected_folder:
            self.label.config(text="请选择文件夹！")
            self.update_status("请选择文件夹！", "danger")
            logging.warning("No folder selected for generation")
            return
        
        # Check for unsupported files
        unsupported_files = check_unsupported_files(self.selected_folder)
        if unsupported_files:
            unsupported_list = "\n".join(unsupported_files[:5])  # Show up to 5 files
            if len(unsupported_files) > 5:
                unsupported_list += "\n... (更多文件未显示)"
            messagebox.showwarning(
                "警告",
                f"检测到不支持的文件类型：\n{unsupported_list}\n"
                "程序只支持图片文件（.png, .jpg, .jpeg），请移除或转换后重试！"
            )
            self.update_status("检测到不支持的文件类型，请检查警告信息！", "danger")
        
        try:
            if messagebox.askyesno("确认重命名", "是否要将所有图片重命名为 1.jpg, 2.jpg 等？此操作不可逆！"):
                self.update_status("正在重命名图片...", "info")
                rename_images(self.selected_folder, self.update_status)
                self.update_status("图片重命名完成！", "success")
            
            structure = collect_files(self.selected_folder)
            if not structure:
                messagebox.showerror("错误", "未找到有效文件或文件夹！请检查文件夹结构。")
                self.update_status("未找到有效文件或文件夹！", "danger")
                return
            output_path = Path(self.selected_folder) / "佐证材料.docx"
            self.progress_bar['value'] = 0
            self.update_status("开始生成文档...", "info")
            success = create_document(structure, self.selected_folder, output_path, 
                                     lambda value: self.progress_bar.__setitem__('value', value) or self.root.update(),
                                     self.update_status)
            if success:
                self.label.config(text=f"文件生成位置: {output_path}")
                messagebox.showinfo("成功", f"文档已生成: {output_path}")
                self.update_status(f"文档已生成: {output_path}", "success")
        except Exception as e:
            logging.error(f"Unexpected error: {e}")
            if "Permission denied" not in str(e):
                messagebox.showerror("错误", f"生成文档失败: {str(e)}")
            self.update_status(f"错误: {str(e)}", "danger")
    
    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = DocGeneratorApp()
    app.run()